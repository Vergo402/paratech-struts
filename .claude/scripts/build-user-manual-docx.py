#!/usr/bin/env python3
"""Regenerate FieldStruts-User-Manual.docx from the canonical docs/USER-MANUAL.md.

Faithful markdown -> docx renderer covering the features actually used in the
manual: title block, section/subsection headings, paragraphs with inline
**bold** / `code` / [links](url), ordered + unordered lists (rendered with
LITERAL numbers/bullets so Word never auto-renumbers across the TOC and the
step lists), pipe tables, and fenced code blocks (the ICS org-chart tree).

It also re-embeds the section screenshots from docs/manual-assets/ (captured
from the live v3.21 app) at the end of each mapped section. Missing images are
skipped gracefully so the script still runs if a shot wasn't captured.

Keeps the .docx content 100% in sync with the markdown source. Run on every
MINOR/MAJOR release after editing docs/USER-MANUAL.md:

    pip install python-docx
    python3 .claude/scripts/build-user-manual-docx.py
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paths are derived from this script's own location (always <repo>/.claude/
# scripts/build-user-manual-docx.py), so the build keeps working no matter where
# the repo lives or how the directory taxonomy changes. Do NOT hardcode an
# absolute repo path here.
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
SRC = os.path.join(ROOT, "docs", "USER-MANUAL.md")
OUT = os.path.join(ROOT, "docs", "FieldStruts-User-Manual.docx")
ASSETS = os.path.join(ROOT, "docs", "manual-assets")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2C, 0x5F, 0x8A)
GREY = RGBColor(0x55, 0x55, 0x55)
CODE_BG = "F2F2F2"

INLINE_RE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')

# Screenshots (docs/manual-assets/) keyed by the EXACT heading text they sit
# under. Each entry: (filename, caption, 'mobile'|'desktop'). Images are flushed
# at the END of the named section (after its text, before the next heading).
MOBILE_W = Inches(2.9)   # phone-portrait shots — narrow + tall
DESKTOP_W = Inches(6.0)  # wide split-view shots

IMAGE_MAP = {
    "Connecting Your Department": [
        ("01-login.png", "The Department ID login screen.", "mobile"),
    ],
    "Using FieldShore on Desktop": [
        ("13-desktop-command.png", "Desktop view — Command tab split layout.", "desktop"),
        ("14-desktop-operations.png", "Desktop view — Operations tab with drilldown sidebar.", "desktop"),
    ],
    "How to Use": [
        ("02-quickfind.png", "Quick Find — entering a measurement and load.", "mobile"),
    ],
    "Results": [
        ("03-quickfind-results.png", "Quick Find results, ranked by safety factor.", "mobile"),
    ],
    "Starting an Operation": [
        ("04-start-operation.png", "The Start New Operation form.", "mobile"),
        ("05-operations-tab.png", "The Operations tab with an active operation.", "mobile"),
    ],
    "Adding a Shore Point": [
        ("06-add-shore-point.png", "The Add Shore Point form.", "mobile"),
    ],
    "Shore Point Lifecycle": [
        ("07-shore-card.png", "A shore point card with its status badge and actions.", "mobile"),
    ],
    "What's on the Command Tab": [
        ("08-command-tab.png", "The Command tab — dashboard tiles and sections.", "mobile"),
    ],
    "Default Roles": [
        ("09-org-chart.png", "The ICS Organization chart (default hierarchy).", "mobile"),
    ],
    "6. Cut Table": [
        ("10-cut-table.png", "The Cut Table view.", "mobile"),
    ],
    "7. Inventory": [
        ("11-inventory.png", "The Inventory tab — apparatus and Quick Add grid.", "mobile"),
    ],
    "8. Settings": [
        ("12-settings.png", "The Settings tab.", "mobile"),
    ],
}


def add_inline(paragraph, text):
    """Render inline markdown (**bold**, `code`, [link](url)) into runs."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = ACCENT
        elif part.startswith('[') and '](' in part:
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            r = paragraph.add_run(m.group(1))
            r.font.color.rgb = ACCENT
            r.underline = True
        else:
            paragraph.add_run(part)


def add_images(doc, entries):
    """Embed a section's screenshots + captions, skipping any missing files."""
    for filename, caption, kind in entries:
        path = os.path.join(ASSETS, filename)
        if not os.path.exists(path):
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run()
        run.add_picture(path, width=(DESKTOP_W if kind == 'desktop' else MOBILE_W))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY


def main():
    with open(SRC, encoding='utf-8') as f:
        lines = f.read().split('\n')

    # ---- title-block metadata ----
    version = last_updated = app_url = ''
    for ln in lines[:12]:
        if ln.startswith('**Version:**'):
            version = ln.split('**Version:**')[1].strip().rstrip('\\').strip()
        elif ln.startswith('**Last updated:**'):
            last_updated = ln.split('**Last updated:**')[1].strip().rstrip('\\').strip()
        elif ln.startswith('**App:**'):
            m = re.search(r'\(([^)]+)\)', ln)
            app_url = m.group(1) if m else ''

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ===== TITLE PAGE =====
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('FieldStruts'); r.font.size = Pt(40); r.bold = True; r.font.color.rgb = NAVY
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('User Manual'); r.font.size = Pt(26); r.font.color.rgb = ACCENT
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Rescue Strut Selection, Inventory Management\n& Shoring Operations for USAR/FEMA')
    r.font.size = Pt(13); r.font.color.rgb = GREY
    for _ in range(2):
        doc.add_paragraph()
    for label in (f'Version {version}', last_updated, app_url):
        if not label:
            continue
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label); r.font.size = Pt(12)
        if label.startswith('Version'):
            r.bold = True; r.font.color.rgb = NAVY
        else:
            r.font.color.rgb = GREY
    doc.add_page_break()

    # ===== BODY =====
    i, n = 0, len(lines)
    while i < n and not lines[i].startswith('## '):
        i += 1

    pending_images = []  # flushed when the next heading arrives or at EOF

    def is_table_row(s):
        return s.strip().startswith('|') and s.strip().endswith('|')

    def emit_heading(level, text):
        nonlocal pending_images
        add_images(doc, pending_images)          # close out previous section
        add_inline(doc.add_heading(level=level), text)
        pending_images = IMAGE_MAP.get(text.strip(), [])

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped in ('---', ''):
            i += 1
            continue

        # fenced code block (ICS org-chart tree)
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run('\n'.join(code_lines))
            r.font.name = 'Consolas'; r.font.size = Pt(9)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), CODE_BG)
            pPr.append(shd)
            continue

        # headings
        if stripped.startswith('#### '):
            emit_heading(3, stripped[5:]); i += 1; continue
        if stripped.startswith('### '):
            emit_heading(2, stripped[4:]); i += 1; continue
        if stripped.startswith('## '):
            emit_heading(1, stripped[3:]); i += 1; continue
        if stripped.startswith('# '):
            emit_heading(0, stripped[2:]); i += 1; continue

        # tables
        if is_table_row(line):
            tbl_rows = []
            while i < n and is_table_row(lines[i]):
                tbl_rows.append(lines[i].strip()); i += 1

            def cells(row):
                return [c.strip() for c in row.strip().strip('|').split('|')]
            header = cells(tbl_rows[0])
            body = [cells(r) for r in tbl_rows[2:]]
            table = doc.add_table(rows=1, cols=len(header))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for j, htext in enumerate(header):
                hdr[j].paragraphs[0].text = ''
                add_inline(hdr[j].paragraphs[0], htext)
                for run in hdr[j].paragraphs[0].runs:
                    run.bold = True
            for brow in body:
                cr = table.add_row().cells
                for j, ctext in enumerate(brow):
                    if j >= len(cr):
                        break
                    cr[j].paragraphs[0].text = ''
                    add_inline(cr[j].paragraphs[0], ctext)
            doc.add_paragraph()
            continue

        # ordered list — LITERAL number (avoids Word cross-list auto-renumber)
        m_ol = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
        if m_ol:
            indent = len(m_ol.group(1))
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Inches(0.3 + (0.3 if indent >= 3 else 0.0))
            pf.first_line_indent = Inches(-0.3)
            pf.space_after = Pt(2)
            p.add_run(f'{m_ol.group(2)}.\t')
            add_inline(p, m_ol.group(3))
            i += 1
            continue

        # unordered list
        m_ul = re.match(r'^(\s*)[-*]\s+(.*)$', line)
        if m_ul:
            indent = len(m_ul.group(1))
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Inches(0.3 + (0.3 if indent >= 2 else 0.0))
            pf.first_line_indent = Inches(-0.2)
            pf.space_after = Pt(2)
            p.add_run('•\t')
            add_inline(p, m_ul.group(2))
            i += 1
            continue

        # plain paragraph
        add_inline(doc.add_paragraph(), stripped)
        i += 1

    add_images(doc, pending_images)  # flush final section's images

    for lvl, color, size in ((1, NAVY, 16), (2, ACCENT, 13), (3, GREY, 11.5)):
        style = doc.styles[f'Heading {lvl}']
        style.font.color.rgb = color
        style.font.size = Pt(size)

    doc.save(OUT)

    embedded = sum(
        1 for entries in IMAGE_MAP.values()
        for fn, _, _ in entries if os.path.exists(os.path.join(ASSETS, fn))
    )
    total = sum(len(v) for v in IMAGE_MAP.values())
    print(f'Wrote {OUT} (Version {version}, {last_updated}) — {embedded}/{total} screenshots embedded')


if __name__ == '__main__':
    main()
